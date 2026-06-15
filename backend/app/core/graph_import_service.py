"""知识图谱导入服务：从 JSON / TXT / Word 文件导入实体与关系。

混合模式（hybrid）：
- JSON：直接按 {entities, relations, synonyms} 结构解析。
- TXT：若多数行可解析为三元组（源 | 关系 | 目标 | 描述）则按结构解析；
        否则视为普通正文，走 LLM 抽取（复用 graph_build_service.extract_graph_data）。
- Word(.docx)：含 ≥3 列的表格→按行解析三元组；否则取正文按 TXT 规则处理。

导入的实体/关系以 source='import' 写入 Neo4j，document_id 为空，行为与手动新增一致，
不随文档删除而清理。实体名沿用放宽后的通用质量过滤（entity_quality）。
"""
import io
import json
import logging

from app.core.entity_quality import clean_entity_record, is_high_quality_entity_name, normalize_entity_name
from app.core.graph_build_service import _guess_entity_type, extract_graph_data
from app.db.neo4j import get_neo4j_driver

logger = logging.getLogger("app")

# 三元组分隔符，按优先级尝试
_TRIPLE_SEPARATORS = ["|", "\t", "→", "-->", "->", "｜"]


def _split_triple_line(line: str) -> list[str] | None:
    """把一行按已知分隔符拆成字段；无可用分隔符时返回 None。"""
    text = line.strip()
    if not text:
        return None
    for sep in _TRIPLE_SEPARATORS:
        if sep in text:
            return [part.strip() for part in text.split(sep) if part.strip() != ""]
    # 退化：中文/英文逗号，仅当恰好切出 3-4 段时才当作三元组
    for sep in ("，", ","):
        if sep in text:
            parts = [part.strip() for part in text.split(sep) if part.strip() != ""]
            if 3 <= len(parts) <= 4:
                return parts
    return None


def _looks_like_triples(lines: list[str]) -> bool:
    nonempty = [ln for ln in lines if ln.strip()]
    if not nonempty:
        return False
    triple_lines = 0
    for ln in nonempty:
        fields = _split_triple_line(ln)
        if fields and len(fields) >= 3:
            triple_lines += 1
    return triple_lines >= max(1, len(nonempty) * 0.5)


def _empty_graph() -> dict:
    return {"entities": [], "relations": [], "synonyms": []}


def _normalize_entity_item(item: dict) -> dict | None:
    """容错读取实体字段（支持 type/entity_type 别名）。"""
    if not isinstance(item, dict):
        return None
    name = str(item.get("name") or item.get("entity") or "").strip()
    if not name:
        return None
    etype = str(item.get("type") or item.get("entity_type") or "").strip()
    desc = str(item.get("description") or item.get("desc") or "").strip()
    props = item.get("properties") if isinstance(item.get("properties"), dict) else {}
    return {"name": name, "type": etype, "description": desc, "properties": props}


def _normalize_relation_item(item) -> dict | None:
    """容错读取关系字段（支持 source/head、target/tail、type/relation 别名，及三元数组）。"""
    if isinstance(item, (list, tuple)):
        if len(item) < 3:
            return None
        source, rel_type, target = str(item[0]).strip(), str(item[1]).strip(), str(item[2]).strip()
        desc = str(item[3]).strip() if len(item) > 3 else ""
    elif isinstance(item, dict):
        source = str(item.get("source") or item.get("head") or item.get("from") or "").strip()
        target = str(item.get("target") or item.get("tail") or item.get("to") or "").strip()
        rel_type = str(item.get("type") or item.get("relation") or item.get("rel") or item.get("relation_type") or "").strip()
        desc = str(item.get("description") or item.get("desc") or "").strip()
    else:
        return None
    if not source or not target:
        return None
    return {"source": source, "target": target, "type": rel_type or "关联", "description": desc}


def _parse_json(raw_text: str) -> dict:
    data = json.loads(raw_text)
    result = _empty_graph()
    if isinstance(data, list):
        # 纯关系三元组数组
        for item in data:
            rel = _normalize_relation_item(item)
            if rel:
                result["relations"].append(rel)
        return result
    if not isinstance(data, dict):
        return result
    for item in data.get("entities") or []:
        ent = _normalize_entity_item(item)
        if ent:
            result["entities"].append(ent)
    for item in data.get("relations") or []:
        rel = _normalize_relation_item(item)
        if rel:
            result["relations"].append(rel)
    for item in data.get("synonyms") or []:
        if isinstance(item, dict):
            original = str(item.get("original") or item.get("name") or "").strip()
            synonym = str(item.get("synonym") or item.get("alias") or "").strip()
            if original and synonym:
                result["synonyms"].append({"original": original, "synonym": synonym})
    return result


def _parse_triple_lines(lines: list[str]) -> dict:
    """把三元组/实体行解析为结构化图谱数据。"""
    result = _empty_graph()
    for ln in lines:
        fields = _split_triple_line(ln)
        if fields is None:
            # 单列：一个实体
            name = ln.strip()
            if name:
                result["entities"].append({"name": name, "type": "", "description": "", "properties": {}})
            continue
        if len(fields) >= 3:
            result["relations"].append({
                "source": fields[0],
                "type": fields[1] or "关联",
                "target": fields[2],
                "description": fields[3] if len(fields) > 3 else "",
            })
        elif len(fields) == 2:
            # 实体 | 类型
            result["entities"].append({"name": fields[0], "type": fields[1], "description": "", "properties": {}})
        elif len(fields) == 1:
            result["entities"].append({"name": fields[0], "type": "", "description": "", "properties": {}})
    return result


def _parse_docx_tables(doc) -> dict:
    """从 Word 表格解析三元组（识别 源/关系/目标/描述 表头，否则按列位置）。"""
    result = _empty_graph()
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        rows = [r for r in rows if any(c for c in r)]
        if len(rows) < 1 or len(rows[0]) < 3:
            continue
        header = rows[0]
        header_joined = "".join(header)
        has_header = any(k in header_joined for k in ("源", "关系", "目标", "实体", "source", "target", "relation"))
        # 列定位
        src_i, rel_i, tgt_i, desc_i = 0, 1, 2, 3
        if has_header:
            for i, h in enumerate(header):
                if any(k in h for k in ("源", "头", "source", "from")):
                    src_i = i
                elif any(k in h for k in ("关系", "类型", "relation", "rel", "type")):
                    rel_i = i
                elif any(k in h for k in ("目标", "尾", "target", "to")):
                    tgt_i = i
                elif any(k in h for k in ("描述", "说明", "desc")):
                    desc_i = i
        data_rows = rows[1:] if has_header else rows
        for r in data_rows:
            if len(r) <= max(src_i, rel_i, tgt_i):
                continue
            source, rel_type, target = r[src_i], r[rel_i], r[tgt_i]
            if not source or not target:
                continue
            desc = r[desc_i] if desc_i < len(r) else ""
            result["relations"].append({
                "source": source,
                "type": rel_type or "关联",
                "target": target,
                "description": desc,
            })
    return result


async def parse_import_file(filename: str, content: bytes) -> dict:
    """根据文件类型与内容自动选择结构化解析或 LLM 抽取。

    返回 {entities, relations, synonyms, mode}，mode ∈ {json, txt, docx-table, llm}。
    """
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()

    # JSON：始终结构化
    if ext == "json":
        try:
            data = _parse_json(content.decode("utf-8-sig", errors="replace"))
            data["mode"] = "json"
            return data
        except Exception as e:
            raise ValueError(f"JSON 解析失败：{str(e)[:200]}")

    # Word
    if ext in ("docx", "doc"):
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        table_data = _parse_docx_tables(doc)
        if table_data["relations"] or table_data["entities"]:
            table_data["mode"] = "docx-table"
            return table_data
        # 无可用表格→取正文按 TXT 规则
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)
        if _looks_like_triples(paragraphs):
            data = _parse_triple_lines(paragraphs)
            data["mode"] = "txt"
            return data
        data = await extract_graph_data(text, doc_label=f"import:{filename}")
        data["mode"] = "llm"
        return data

    # TXT / MD / 其它纯文本
    text = content.decode("utf-8-sig", errors="replace")
    lines = text.splitlines()
    if _looks_like_triples(lines):
        data = _parse_triple_lines(lines)
        data["mode"] = "txt"
        return data
    data = await extract_graph_data(text, doc_label=f"import:{filename}")
    data["mode"] = "llm"
    return data


async def _write_import_graph(parsed: dict) -> dict:
    """把解析出的实体/关系/同义词以 source='import' 写入 Neo4j。"""
    skipped = 0

    # 1. 清洗实体（放宽后的通用质量过滤）
    entities: list[dict] = []
    entity_names: set[str] = set()
    for raw in parsed.get("entities") or []:
        cleaned = clean_entity_record(raw)
        if not cleaned:
            skipped += 1
            continue
        if cleaned["name"] in entity_names:
            continue
        entity_names.add(cleaned["name"])
        entities.append(cleaned)

    # 2. 关系端点若不在实体集合中则自动补建实体
    relations: list[dict] = []
    for rel in parsed.get("relations") or []:
        source = normalize_entity_name(str(rel.get("source", "")))
        target = normalize_entity_name(str(rel.get("target", "")))
        if not is_high_quality_entity_name(source) or not is_high_quality_entity_name(target):
            skipped += 1
            continue
        for endpoint in (source, target):
            if endpoint not in entity_names:
                entity_names.add(endpoint)
                entities.append({"name": endpoint, "type": _guess_entity_type(endpoint), "description": ""})
        relations.append({**rel, "source": source, "target": target})

    # 3. 同义词
    synonyms: list[dict] = []
    for syn in parsed.get("synonyms") or []:
        original = normalize_entity_name(str(syn.get("original", "")))
        synonym = normalize_entity_name(str(syn.get("synonym", "")))
        if original and synonym and synonym != original:
            synonyms.append({"original": original, "synonym": synonym})

    if not entities and not relations:
        return {"entities": 0, "relations": 0, "synonyms": 0, "skipped": skipped}

    driver = await get_neo4j_driver()

    for entity in entities:
        props = entity.get("properties") or {}
        props_json = json.dumps(props, ensure_ascii=False) if isinstance(props, dict) else "{}"
        try:
            async with driver.session() as session:
                await session.run(
                    "MERGE (e:Entity {name: $name}) "
                    "ON CREATE SET e.source = 'import' "
                    "SET e.type = $type, "
                    "    e.description = CASE WHEN $desc <> '' THEN $desc ELSE e.description END, "
                    "    e.properties = $props",
                    name=entity["name"],
                    type=entity.get("type") or "导入",
                    desc=entity.get("description", "") or "",
                    props=props_json,
                )
        except Exception as exc:
            logger.warning(f"导入实体 '{entity['name']}' 失败: {str(exc)[:200]}")

    written_relations = 0
    for rel in relations:
        try:
            async with driver.session() as session:
                result = await session.run(
                    "MATCH (a:Entity {name: $src}), (b:Entity {name: $tgt}) "
                    "MERGE (a)-[r:RELATES {rel_type: $rel_type}]->(b) "
                    "SET r.description = $desc, r.source = 'import' "
                    "RETURN r",
                    src=rel["source"], tgt=rel["target"],
                    rel_type=rel.get("type", "关联") or "关联",
                    desc=rel.get("description", "") or "",
                )
                if await result.single():
                    written_relations += 1
        except Exception as exc:
            logger.warning(f"导入关系 '{rel['source']}→{rel['target']}' 失败: {str(exc)[:200]}")

    for syn in synonyms:
        try:
            async with driver.session() as session:
                await session.run(
                    "MERGE (s:Synonym {original: $orig, synonym: $syn})",
                    orig=syn["original"], syn=syn["synonym"],
                )
        except Exception as exc:
            logger.warning(f"导入同义词失败: {str(exc)[:200]}")

    return {
        "entities": len(entities),
        "relations": written_relations,
        "synonyms": len(synonyms),
        "skipped": skipped,
    }


async def import_graph_file(filename: str, content: bytes) -> dict:
    """对外入口：解析文件并写入图谱，返回导入汇总。"""
    if not content:
        return {"success": False, "message": "文件为空"}
    try:
        parsed = await parse_import_file(filename, content)
    except ValueError as e:
        return {"success": False, "message": str(e)}
    except Exception as e:
        logger.error(f"导入文件解析失败 {filename}: {e}", exc_info=True)
        return {"success": False, "message": f"文件解析失败：{str(e)[:200]}"}

    mode = parsed.pop("mode", "?")
    summary = await _write_import_graph(parsed)
    summary.update({
        "success": True,
        "mode": mode,
        "message": f"导入完成：实体 {summary['entities']} 个，关系 {summary['relations']} 条"
                   + (f"，跳过 {summary['skipped']} 项" if summary.get("skipped") else ""),
    })
    return summary
