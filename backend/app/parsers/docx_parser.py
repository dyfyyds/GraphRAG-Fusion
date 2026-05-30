from docx import Document as DocxDocument
from app.parsers.base import BaseParser, ParsedContent


class DocxParser(BaseParser):
    def parse(self, file_path: str) -> ParsedContent:
        content = ParsedContent()
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        content.text = "\n".join(full_text)
        content.page_map.append({"page": 1, "text": content.text})

        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            content.tables.append(rows)
        return content
